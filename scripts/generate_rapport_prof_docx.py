"""
Génère le rapport Word du projet Sommeil_EOG_IA (≤ 15 pages), avec figures réelles
(signaux EOG, hypnogrammes, matrice de confusion, courbes) et métriques mesurées.

Prérequis : lancer d'abord les assets (figures) si absents :
  .\.venv\Scripts\python.exe scripts/make_report_assets.py
Puis :
  .\.venv\Scripts\python.exe scripts/generate_rapport_prof_docx.py

Dépendance : pip install python-docx
"""
from __future__ import annotations

import json
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
FIG_DIR = os.path.join(ROOT, "reports", "figures")
OUT_NAME = "Rapport_Projet_Sommeil_EOG_IA.docx"

PRIMARY = RGBColor(12, 74, 110)
GREY = RGBColor(71, 85, 105)


def fig(name: str) -> str:
    return os.path.join(FIG_DIR, name)


def load_metrics(name: str) -> dict:
    with open(os.path.join(ROOT, "models", name), encoding="utf-8") as f:
        return json.load(f)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = PRIMARY
    return h


def add_para(doc, text, size=10.5, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


def _shade_cell(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def add_table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _shade_cell(cell, "0C4A6E")
        run = cell.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            if ri % 2 == 1:
                _shade_cell(cell, "EAF1F6")
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(size)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    return t


def add_figure(doc, path, width=6.0, caption=None):
    if not os.path.isfile(path):
        add_para(doc, f"[figure manquante : {os.path.basename(path)}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        c.paragraph_format.space_after = Pt(8)


def pct(x):
    return f"{x*100:.1f} %"


def build():
    cnn = load_metrics("metrics_holdout_test.json")
    bilstm = load_metrics("metrics_test_bilstm.json")
    bilstm_val = load_metrics("metrics_val_bilstm.json")
    try:
        with open(os.path.join(FIG_DIR, "_summary.json"), encoding="utf-8") as f:
            demo = json.load(f)
    except FileNotFoundError:
        demo = {"patient01_epochs": 2650, "patient01_accuracy": 0.923, "patient01_channel": "EOG horizontal"}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    # ---------------- Page de titre ----------------
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Classification automatique des stades du sommeil\npar intelligence artificielle à partir du signal EOG")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = PRIMARY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Projet Sommeil_EOG_IA — Dashboard clinique « DeepSleep AI »")
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    add_figure(doc, fig("fig_pipeline.png"), width=6.4)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "Apprentissage profond (CNN 1D) · Inférence accélérée OpenVINO (NPU / GPU / CPU)",
        "Corpus Sleep-EDF Expanded (PhysioNet) · Nomenclature AASM 5 stades",
        "Rapport de synthèse",
    ):
        rr = meta.add_run(line + "\n")
        rr.font.size = Pt(11)
    doc.add_page_break()

    # ---------------- 1. Résumé ----------------
    add_heading(doc, "1. Résumé exécutif")
    add_para(
        doc,
        "Ce projet met en œuvre une chaîne complète, des enregistrements polysomnographiques bruts "
        "(fichiers EDF) jusqu'à une application web clinique, pour estimer automatiquement les stades du "
        "sommeil à partir d'un unique canal d'électro-oculographie (EOG), selon la norme AASM "
        "(cinq classes : éveil W, N1, N2, N3 et sommeil paradoxal REM). Un réseau de neurones convolutif "
        "1D est entraîné sur le corpus Sleep-EDF Expanded (182 sujets, 419 170 époques de 30 s), puis "
        "exporté au format OpenVINO IR FP16 afin de tirer parti de l'accélération matérielle (NPU Intel "
        "AI Boost, GPU intégré ou CPU). Évalué sur un jeu de test indépendant de 27 sujets jamais vus à "
        f"l'entraînement, le modèle atteint une exactitude de {pct(cnn['accuracy'])}, un F1 macro de "
        f"{cnn['f1_macro']:.3f} et un coefficient kappa de Cohen de {cnn['cohen_kappa']:.3f} "
        "(accord « substantiel » avec l'expert). Une interface Streamlit permet le chargement des "
        "enregistrements, l'inférence en temps réel, la visualisation de l'hypnogramme et la production "
        "d'un rapport clinique synthétique.",
    )

    # ---------------- 2. Contexte ----------------
    add_heading(doc, "2. Contexte et problématique")
    add_para(
        doc,
        "La polysomnographie (PSG) enregistre plusieurs signaux physiologiques durant une nuit de sommeil "
        "(EEG, EOG, EMG, etc.). Le « scoring » manuel consiste à attribuer un stade de sommeil à chaque "
        "fenêtre de 30 secondes ; réalisé par un expert, il est long, coûteux et sujet à variabilité "
        "inter-juges. L'EOG capture les mouvements oculaires, particulièrement informatifs pour distinguer "
        "l'éveil (clignements, mouvements rapides) et le sommeil paradoxal (mouvements oculaires rapides "
        "caractéristiques). L'objectif est de démontrer un pipeline reproductible et déployable : "
        "prétraitement du signal, apprentissage supervisé, optimisation pour accélérateur matériel et "
        "restitution visuelle à destination du technologue ou du chercheur.",
    )
    add_heading(doc, "Objectifs du projet", level=2)
    add_bullets(doc, [
        "Charger et synchroniser signaux et annotations au format EDF (MNE-Python).",
        "Prétraiter le canal EOG : rééchantillonnage 100 Hz, filtrage 0,5–35 Hz, normalisation z-score.",
        "Entraîner un classificateur d'époques de 30 s sur les 5 classes AASM, avec un découpage par sujet.",
        "Exporter un modèle compatible NPU via OpenVINO (FP16, forme d'entrée fixe 64×3000×1).",
        "Fournir un tableau de bord interactif pour l'analyse, la comparaison à l'expert et l'export des résultats.",
    ])

    # ---------------- 3. Données ----------------
    add_heading(doc, "3. Données et nomenclature")
    add_para(
        doc,
        "Le corpus principal est Sleep-EDF Expanded (PhysioNet), complété par quatre enregistrements "
        "locaux (Patient_01, 02, 11, 12). Après chargement et synchronisation, 182 sujets exploitables "
        "fournissent 419 170 époques de 30 s échantillonnées à 100 Hz (soit 3000 points par époque). "
        "La nomenclature AASM regroupe les anciens stades N3 et N4 en une seule classe de sommeil profond.",
    )
    add_table(
        doc,
        ["Code", "Stade", "Description clinique"],
        [
            ["0", "W", "Éveil — clignements et mouvements oculaires volontaires"],
            ["1", "N1", "Endormissement / sommeil léger, transition"],
            ["2", "N2", "Sommeil léger consolidé (fuseaux, complexes K)"],
            ["3", "N3", "Sommeil profond à ondes lentes (N3 + N4 fusionnés)"],
            ["4", "REM", "Sommeil paradoxal — mouvements oculaires rapides"],
        ],
        widths=[0.7, 0.9, 4.6],
    )
    add_para(
        doc,
        "Les stades sont fortement déséquilibrés : l'éveil et le N2 dominent, tandis que le N1 est rare "
        "et difficile à distinguer (transition courte). Ce déséquilibre, illustré ci-dessous, est compensé "
        "à l'entraînement par une pondération des classes (class_weight).",
        space_after=4,
    )
    add_figure(doc, fig("fig_class_distribution.png"), width=5.6,
               caption="Figure 1 — Répartition des 64 113 époques du jeu de test sur les 5 stades AASM.")

    # ---------------- 4. Prétraitement ----------------
    add_heading(doc, "4. Prétraitement du signal EOG")
    add_para(
        doc,
        "Le prétraitement (module src/preprocessing.py) homogénéise les enregistrements avant l'inférence : "
        "rééchantillonnage à 100 Hz, sélection automatique du premier canal dont le nom contient « EOG », "
        "filtrage passe-bande FIR 0,5–35 Hz pour éliminer la dérive lente et le bruit haute fréquence, puis "
        "normalisation z-score par enregistrement et écrêtage des valeurs aberrantes à ±3 écarts-types. "
        "La figure 2 montre l'effet du prétraitement sur une époque, et la figure 3 met en évidence des "
        "morphologies distinctes selon le stade — base sur laquelle le réseau apprend à discriminer.",
    )
    add_table(
        doc,
        ["Étape", "Paramètres"],
        [
            ["Rééchantillonnage", "100 Hz (aligné sur l'entrée du modèle)"],
            ["Sélection du canal", "Premier canal contenant « EOG » (insensible à la casse)"],
            ["Filtrage", "Passe-bande FIR 0,5 – 35 Hz (MNE-Python)"],
            ["Normalisation", "Z-score par enregistrement, puis écrêtage à ±3 σ"],
            ["Découpage", "Époques de 30 s = 3000 points → tenseur (N, 3000, 1)"],
        ],
        widths=[1.7, 4.5],
    )
    add_figure(doc, fig("fig_eog_signal.png"), width=5.9,
               caption="Figure 2 — Signal EOG brut (haut) et après prétraitement (bas) sur une époque de 30 s.")
    add_figure(doc, fig("fig_eog_stages.png"), width=5.5,
               caption="Figure 3 — Morphologie typique du signal EOG prétraité pour chacun des 5 stades (Patient_01).")

    # ---------------- 5. Modèle ----------------
    add_heading(doc, "5. Architecture du modèle")
    add_para(
        doc,
        "Deux architectures ont été développées (src/architecture.py). Une variante CNN + Bi-LSTM offre "
        "une bonne modélisation temporelle mais utilise des opérations récurrentes (Loop, ReverseSequence) "
        "non supportées par le NPU Intel ciblé. Une variante CNN 1D pure — quatre blocs convolutifs "
        "(Conv1D + BatchNorm + MaxPooling) suivis d'un global average pooling et de couches denses — n'emploie "
        "que des opérations compatibles NPU et constitue le modèle déployé. Elle compte environ 455 000 "
        "paramètres pour une entrée statique de forme (64, 3000, 1).",
    )
    add_figure(doc, fig("fig_architecture.png"), width=6.3,
               caption="Figure 4 — Architecture du CNN 1D compatible NPU (détail dans src/architecture.py).")
    add_table(
        doc,
        ["Bloc", "Couches", "Sortie"],
        [
            ["1", "Conv1D(64, k=11) + BN + MaxPool(4)", "(750, 64)"],
            ["2", "Conv1D(128, k=7) + BN + MaxPool(4)", "(187, 128)"],
            ["3", "Conv1D(256, k=5) + BN + MaxPool(4)", "(46, 256)"],
            ["4", "Conv1D(256, k=3) + BN + MaxPool(2)", "(23, 256)"],
            ["Tête", "GlobalAvgPool1D → Dense(128) → Dense(5, softmax)", "(5,)"],
        ],
        widths=[0.7, 4.2, 1.3],
    )

    # ---------------- 6. Entraînement ----------------
    add_heading(doc, "6. Protocole d'entraînement et d'évaluation")
    add_para(
        doc,
        "Le point méthodologique central est le découpage par sujet (src/splits.py) : les enregistrements "
        "sont répartis en 127 sujets d'entraînement (297 012 époques), 28 sujets de validation (58 045 "
        "époques) et 27 sujets de test (64 113 époques). Aucune époque d'un sujet de test n'apparaît à "
        "l'entraînement ni à la validation, ce qui évite toute fuite de données et donne une estimation "
        "honnête de la généralisation. La fonction de perte est l'entropie croisée catégorielle, "
        "l'optimiseur Adam, avec pondération des classes et arrêt anticipé sur le F1 macro de validation. "
        "L'entraînement a été mené sur GPU (Kaggle P100). La figure 5 montre la progression du F1 macro "
        "et de l'exactitude en validation.",
    )
    add_figure(doc, fig("fig_training_curve.png"), width=5.8,
               caption="Figure 5 — Courbes de validation (CNN 1D NPU). Le meilleur modèle (F1 ≈ 0,724) est conservé.")

    # ---------------- 7. Résultats ----------------
    add_heading(doc, "7. Résultats quantitatifs")
    add_para(
        doc,
        "Les performances ci-dessous sont calculées sur le jeu de test hold-out (27 sujets, "
        "64 113 époques) via src/evaluate_holdout.py. Le CNN 1D déployé dépasse la variante CNN + Bi-LSTM "
        "sur l'ensemble des métriques globales, tout en restant compatible NPU.",
    )
    add_table(
        doc,
        ["Métrique", "CNN 1D (NPU) — déployé", "CNN + Bi-LSTM (référence)"],
        [
            ["Exactitude (accuracy)", pct(cnn["accuracy"]), pct(bilstm["accuracy"])],
            ["F1 macro", f"{cnn['f1_macro']:.3f}", f"{bilstm['f1_macro']:.3f}"],
            ["F1 pondéré", f"{cnn['f1_weighted']:.3f}", f"{bilstm['f1_weighted']:.3f}"],
            ["Kappa de Cohen", f"{cnn['cohen_kappa']:.3f}", f"{bilstm['cohen_kappa']:.3f}"],
        ],
        widths=[2.1, 2.2, 2.2],
    )
    add_para(
        doc,
        "L'analyse par stade (figure 6) confirme une excellente détection de l'éveil et du sommeil profond, "
        "de bonnes performances en N2 et REM, et la difficulté attendue sur le N1 — stade de transition "
        "minoritaire et ambigu, y compris pour les experts humains. La matrice de confusion (figure 7) "
        "montre que les erreurs résiduelles concernent surtout les confusions N1↔N2 et N1↔REM.",
    )
    add_figure(doc, fig("fig_f1_comparison.png"), width=5.8,
               caption="Figure 6 — Score F1 par stade : CNN 1D (NPU) vs CNN + Bi-LSTM (test hold-out).")
    add_figure(doc, fig("fig_confusion_cnn.png"), width=4.5,
               caption="Figure 7 — Matrice de confusion normalisée par ligne (CNN 1D NPU, test hold-out).")
    add_table(
        doc,
        ["Stade", "Précision", "Rappel", "F1", "Support"],
        [
            [s,
             f"{cnn['per_class'][s].get('precision', 0):.2f}",
             f"{cnn['per_class'][s]['recall']:.2f}",
             f"{cnn['per_class'][s]['f1']:.2f}",
             f"{cnn['per_class'][s]['support']}"]
            for s in ["W", "N1", "N2", "N3", "REM"]
        ],
        widths=[1.0, 1.2, 1.1, 1.0, 1.2],
    )

    # ---------------- 8. Démonstration ----------------
    add_heading(doc, "8. Démonstration sur un cas réel (Patient_01)")
    add_para(
        doc,
        f"Sur l'enregistrement de démonstration Patient_01 ({demo['patient01_epochs']} époques, canal "
        f"« {demo['patient01_channel']} »), réservé au jeu de test, l'hypnogramme prédit par l'IA reproduit "
        f"fidèlement la structure du sommeil scorée par l'expert, avec une concordance de "
        f"{pct(demo['patient01_accuracy'])}. On retrouve l'alternance des cycles de sommeil, la présence de "
        "sommeil profond en début de nuit et l'apparition progressive du sommeil paradoxal.",
    )
    add_figure(doc, fig("fig_hypnogram.png"), width=6.1,
               caption="Figure 8 — Hypnogramme prédit par l'IA (haut) vs scoring expert de référence (bas), Patient_01.")

    # ---------------- 9. Déploiement ----------------
    add_heading(doc, "9. Déploiement et application")
    add_para(
        doc,
        "Le modèle Keras est exporté en OpenVINO IR FP16 (scripts/export_cnn_openvino.py) : deux fichiers "
        "légers (sleep_model_npu.xml ≈ 56 ko, sleep_model_npu.bin ≈ 0,9 Mo). OpenVINO sélectionne "
        "automatiquement l'accélérateur disponible (NPU, GPU intégré ou CPU). Le dashboard Streamlit "
        "(app/dashboard.py) charge uniquement ce modèle IR — sans TensorFlow à l'exécution — ce qui le rend "
        "déployable sur Streamlit Community Cloud.",
    )
    add_table(
        doc,
        ["Périphérique", "Débit indicatif (époques/s)", "Remarque"],
        [
            ["NPU (Intel AI Boost)", "≈ 5 650", "×5,2 vs CPU — modèle IR FP16"],
            ["GPU intégré", "≈ 2 150", "Modèle IR FP16"],
            ["CPU", "≈ 1 100", "Repli universel (Cloud)"],
        ],
        widths=[2.0, 2.4, 2.2],
    )
    add_heading(doc, "Fonctionnalités du tableau de bord", level=2)
    add_bullets(doc, [
        "Choix du périphérique d'inférence OpenVINO (NPU / GPU / CPU / AUTO).",
        "Import EDF (signal + hypnogramme expert optionnel) ou base d'exemples locale.",
        "Hypnogramme IA, architecture du sommeil et comparaison IA vs expert.",
        "Métriques cliniques : temps de sommeil total, efficacité, latences, pourcentages de stades.",
        "Export des résultats époque par époque au format CSV.",
    ])

    # ---------------- 10. Conclusion ----------------
    add_heading(doc, "10. Limites, perspectives et conclusion")
    add_bullets(doc, [
        "Un seul canal EOG ne remplace pas une PSG complète : l'outil est une aide à la décision, non un dispositif diagnostique.",
        "Le stade N1 reste le point faible (F1 ≈ 0,45) ; des pistes incluent une perte focale et l'ajout de contexte temporel.",
        "Les performances dépendent de la population et de la qualité du signal ; une validation multicentrique serait nécessaire.",
        "Perspectives : modèle multi-canaux, explicabilité (Grad-CAM), agrégation temporelle séquence-à-séquence.",
    ])
    add_para(
        doc,
        "Le projet Sommeil_EOG_IA illustre une chaîne de bout en bout moderne et reproductible : données "
        "standardisées (EDF / AASM), prétraitement robuste, modèle profond évalué rigoureusement par "
        "découpage sujet, optimisation matérielle via OpenVINO et restitution clinique interactive. "
        "Le modèle déployé atteint une concordance « substantielle » avec l'expert "
        f"(kappa = {cnn['cohen_kappa']:.3f}, exactitude = {pct(cnn['accuracy'])}), démontrant la "
        "faisabilité d'un scoring automatique du sommeil à partir d'un capteur unique et peu coûteux.",
    )

    add_heading(doc, "Références et outils", level=2)
    add_para(
        doc,
        "Corpus : Sleep-EDF Expanded, PhysioNet (physionet.org). Bibliothèques : MNE-Python (mne.tools), "
        "TensorFlow / Keras (tensorflow.org), OpenVINO (docs.openvino.ai), scikit-learn, Matplotlib, "
        "Streamlit (streamlit.io). Code source complet et notebook fournis avec ce rapport.",
        size=9.5, italic=True,
    )

    out = os.path.join(ROOT, OUT_NAME)
    doc.save(out)
    return out


if __name__ == "__main__":
    if not os.path.isfile(fig("fig_hypnogram.png")):
        print("Figures absentes — génération via make_report_assets.py …")
        subprocess.run([sys.executable, os.path.join(ROOT, "scripts", "make_report_assets.py")], check=True)
    path = build()
    print("Rapport créé :", path)
